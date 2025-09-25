#!/usr/bin/env python3
"""
Python SOL Document Processor
Processes .docx/.doc files containing SOL standards and converts them to structured data
"""
import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime
from dataclasses import dataclass, asdict
from dataclasses_json import dataclass_json

import docx
from docx import Document
import openai
from sqlalchemy import create_engine, or_
from sqlalchemy.orm import sessionmaker
from pydantic import BaseModel, ValidationError

from server.models import Base, SolStandard

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass_json
@dataclass
class SubObjective:
    """Sub-objective within a SOL standard"""
    code: str
    description: str


@dataclass_json
@dataclass
class SOLStandardData:
    """Individual SOL Standard with all metadata"""
    standard_code: str
    subject: str
    grade: str
    strand: str
    title: Optional[str] = None
    description: str = ""
    sub_objectives: List[SubObjective] = None
    prerequisites: List[str] = None
    connections: List[str] = None
    key_terms: List[str] = None
    difficulty: str = "grade-level"  # foundational, grade-level, advanced
    cognitive_complexity: str = "skill"  # recall, skill, strategic, extended

    def __post_init__(self):
        if self.sub_objectives is None:
            self.sub_objectives = []
        if self.prerequisites is None:
            self.prerequisites = []
        if self.connections is None:
            self.connections = []
        if self.key_terms is None:
            self.key_terms = []


@dataclass_json
@dataclass
class SOLDocumentData:
    """Complete processed SOL document"""
    metadata: Dict[str, Any]
    standards: List[SOLStandardData]


class SOLProcessor:
    """Main SOL document processing class"""

    def __init__(self, database_url: Optional[str] = None, openai_api_key: Optional[str] = None):
        """Initialize the SOL processor

        Args:
            database_url: PostgreSQL connection string
            openai_api_key: OpenAI API key for processing
        """
        # Set up OpenAI
        self.openai_api_key = openai_api_key or os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("OpenAI API key is required. Set OPENAI_API_KEY environment variable.")

        # Set up database
        self.database_url = database_url or os.getenv('DATABASE_URL')
        if not self.database_url:
            raise ValueError("Database URL is required. Set DATABASE_URL environment variable.")

        self.engine = create_engine(self.database_url)
        self.SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=self.engine)

        # Create tables if they don't exist
        Base.metadata.create_all(bind=self.engine)

        logger.info("SOL Processor initialized successfully")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract raw text from a .docx file

        Args:
            file_path: Path to the .docx file

        Returns:
            Raw text content
        """
        try:
            doc = Document(file_path)
            full_text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            return '\n'.join(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def create_sol_extraction_prompt(self, document_text: str, file_name: str) -> str:
        """Create the GPT-4 prompt for SOL extraction"""
        return f"""
Extract Virginia Standards of Learning (SOL) data from this document and return it as structured JSON.

Document: {file_name}
Content:
{document_text}

Please extract ALL standards from this document and format them according to this JSON schema:

{{
  "metadata": {{
    "document_title": "string - title of the document",
    "subject": "string - subject (mathematics, science, english, etc.)",
    "grade_level": "string - grade level (K, 1, 2, 3, etc. or course name)",
    "year_approved": "string - year if mentioned",
    "total_standards": "number - count of main standards"
  }},
  "standards": [
    {{
      "standard_code": "string - SOL code (e.g., '3.NS.1', 'ALG.A.1')",
      "subject": "string - subject area",
      "grade": "string - grade level",
      "strand": "string - content strand/domain",
      "title": "string - short title (optional)",
      "description": "string - full description",
      "sub_objectives": [
        {{
          "code": "string - sub-objective code (e.g., '3.NS.1.a')",
          "description": "string - sub-objective description"
        }}
      ],
      "prerequisites": ["array of related lower-grade standard codes"],
      "connections": ["array of related same/higher-grade standard codes"],
      "key_terms": ["array of important vocabulary"],
      "difficulty": "foundational|grade-level|advanced",
      "cognitive_complexity": "recall|skill|strategic|extended"
    }}
  ]
}}

IMPORTANT EXTRACTION RULES:
1. Extract EVERY standard mentioned in the document
2. Identify the correct standard code format (e.g., grade.strand.number)
3. Group sub-objectives under their main standard
4. Infer subject and grade from document context
5. Map content to appropriate strands (Number Sense, Geometry, etc.)
6. Extract key vocabulary and terms mentioned
7. Identify cognitive complexity based on verbs used (remember=recall, apply=skill, analyze=strategic, create=extended)
8. Return ONLY valid JSON - no markdown, no explanations, no additional text

CRITICAL: Your response must be pure JSON that starts with {{ and ends with }}. Do not include ```json or any other formatting.
"""

    def process_sol_document(self, document_text: str, file_name: str) -> SOLDocumentData:
        """Process SOL document text with GPT-4 to extract structured data

        Args:
            document_text: Raw text from the document
            file_name: Name of the source file

        Returns:
            Structured SOL document data
        """
        prompt = self.create_sol_extraction_prompt(document_text, file_name)

        try:
            from openai import OpenAI
            client = OpenAI(api_key=self.openai_api_key)

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert in Virginia Standards of Learning (SOL) who extracts structured data from educational documents. You must return valid JSON that matches the provided schema exactly."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_tokens=4000
            )

            response_text = response.choices[0].message.content
            logger.info(f"GPT-4 response received: {len(response_text)} characters")

            # Clean and parse JSON response
            if not response_text or response_text.strip() == "":
                raise ValueError("Empty response from OpenAI")

            # Remove any markdown formatting or extra text
            response_text = response_text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()

            # Debug log the response
            logger.debug(f"Cleaned response: {response_text[:200]}...")

            try:
                parsed_data = json.loads(response_text)
            except json.JSONDecodeError as e:
                logger.error(f"JSON parsing failed. Response was: {response_text[:500]}...")
                raise ValueError(f"Invalid JSON response from OpenAI: {e}")

            # Convert to our data structures
            standards_list = []
            for std_data in parsed_data.get('standards', []):

                # Convert sub_objectives
                sub_objectives = []
                for sub_obj in std_data.get('sub_objectives', []):
                    sub_objectives.append(SubObjective(
                        code=sub_obj.get('code', ''),
                        description=sub_obj.get('description', '')
                    ))

                standard = SOLStandardData(
                    standard_code=std_data.get('standard_code', ''),
                    subject=std_data.get('subject', '').lower(),
                    grade=std_data.get('grade', ''),
                    strand=std_data.get('strand', ''),
                    title=std_data.get('title'),
                    description=std_data.get('description', ''),
                    sub_objectives=sub_objectives,
                    prerequisites=std_data.get('prerequisites', []),
                    connections=std_data.get('connections', []),
                    key_terms=std_data.get('key_terms', []),
                    difficulty=std_data.get('difficulty', 'grade-level'),
                    cognitive_complexity=std_data.get('cognitive_complexity', 'skill')
                )
                standards_list.append(standard)

            return SOLDocumentData(
                metadata=parsed_data.get('metadata', {}),
                standards=standards_list
            )

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON response: {e}")
            raise
        except Exception as e:
            logger.error(f"Error processing document with GPT-4: {e}")
            raise

    def save_to_database(self, sol_data: SOLDocumentData) -> int:
        """Save processed SOL data to PostgreSQL database

        Args:
            sol_data: Processed SOL document data

        Returns:
            Number of standards saved
        """
        session = self.SessionLocal()
        saved_count = 0

        try:
            for standard_data in sol_data.standards:

                # Create unique ID
                standard_id = f"{standard_data.subject}-{standard_data.grade}-{standard_data.standard_code}"

                # Check if standard already exists
                existing = session.query(SolStandard).filter(SolStandard.id == standard_id).first()

                # Prepare metadata
                metadata = {
                    'standard_code': standard_data.standard_code,
                    'title': standard_data.title,
                    'sub_objectives': [asdict(sub) for sub in standard_data.sub_objectives],
                    'prerequisites': standard_data.prerequisites,
                    'connections': standard_data.connections,
                    'key_terms': standard_data.key_terms,
                    'difficulty': standard_data.difficulty,
                    'cognitive_complexity': standard_data.cognitive_complexity,
                    'processed_from': sol_data.metadata.get('document_title', 'Unknown'),
                    'processed_at': datetime.now().isoformat()
                }

                if existing:
                    # Update existing standard
                    existing.standard_code = standard_data.standard_code
                    existing.subject = standard_data.subject
                    existing.grade = standard_data.grade
                    existing.strand = standard_data.strand
                    existing.title = standard_data.title
                    existing.description = standard_data.description
                    existing.sol_metadata = metadata
                    logger.info(f"Updated existing standard: {standard_id}")
                else:
                    # Create new standard
                    new_standard = SolStandard(
                        id=standard_id,
                        standard_code=standard_data.standard_code,
                        subject=standard_data.subject,
                        grade=standard_data.grade,
                        strand=standard_data.strand,
                        title=standard_data.title,
                        description=standard_data.description,
                        sol_metadata=metadata
                    )
                    session.add(new_standard)
                    logger.info(f"Created new standard: {standard_id}")

                saved_count += 1

            session.commit()
            logger.info(f"Successfully saved {saved_count} standards to database")
            return saved_count

        except Exception as e:
            session.rollback()
            logger.error(f"Error saving to database: {e}")
            raise
        finally:
            session.close()

    def process_file(self, file_path: str, output_path: Optional[str] = None) -> SOLDocumentData:
        """Process a single SOL document file end-to-end

        Args:
            file_path: Path to the .docx file
            output_path: Optional path to save JSON output

        Returns:
            Processed SOL document data
        """
        logger.info(f"🔄 Processing SOL document: {file_path}")

        # Extract text from document
        document_text = self.extract_text_from_docx(file_path)
        logger.info(f"📄 Extracted {len(document_text)} characters from document")

        # Process with GPT-4
        sol_data = self.process_sol_document(document_text, file_path)
        logger.info(f"🧠 Extracted {len(sol_data.standards)} standards")

        # Optionally save structured output to file
        if output_path:
            with open(output_path, 'w') as f:
                json.dump(asdict(sol_data), f, indent=2)
            logger.info(f"💾 Saved structured data to {output_path}")

        # Save to database
        saved_count = self.save_to_database(sol_data)
        logger.info(f"✅ Saved {saved_count} standards to database")

        return sol_data

    def process_directory(self, directory_path: str, pattern: str = "*.docx") -> List[SOLDocumentData]:
        """Process multiple SOL documents in a directory

        Args:
            directory_path: Path to directory containing SOL documents
            pattern: File pattern to match (default: *.docx)

        Returns:
            List of processed SOL document data
        """
        directory = Path(directory_path)
        files = list(directory.glob(pattern))

        logger.info(f"📁 Found {len(files)} documents to process")

        results = []
        for file_path in files:
            try:
                result = self.process_file(str(file_path))
                results.append(result)

                # Small delay to avoid rate limits
                import time
                time.sleep(1)

            except Exception as e:
                logger.error(f"❌ Failed to process {file_path}: {e}")
                continue

        return results

    def get_database_stats(self) -> Dict[str, Any]:
        """Get statistics about SOL data in database"""
        session = self.SessionLocal()

        try:
            total_standards = session.query(SolStandard).count()

            # Group by subject and grade
            standards = session.query(SolStandard).all()

            by_subject = {}
            for std in standards:
                if std.subject not in by_subject:
                    by_subject[std.subject] = {}
                if std.grade not in by_subject[std.subject]:
                    by_subject[std.subject][std.grade] = 0
                by_subject[std.subject][std.grade] += 1

            return {
                'total_standards': total_standards,
                'by_subject_grade': by_subject
            }

        finally:
            session.close()


if __name__ == "__main__":
    # Simple CLI usage
    import sys

    if len(sys.argv) < 2:
        print("Usage: python sol_processor.py <path-to-docx-file> [output-json-path]")
        sys.exit(1)

    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        processor = SOLProcessor()
        result = processor.process_file(file_path, output_path)
        print("✅ SOL document processing completed successfully")

    except Exception as e:
        print(f"❌ Processing failed: {e}")
        sys.exit(1)