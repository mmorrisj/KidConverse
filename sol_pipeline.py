#!/usr/bin/env python3
"""
Two-Stage SOL Document Processing Pipeline
Stage 1: Detection - Determine if document contains SOL content
Stage 2: Transformation - Extract and transform SOL content to standards.json format
"""

import os
import json
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field

import docx
from docx import Document
import openpyxl
from openpyxl import load_workbook
from openai import OpenAI

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class SOLDetectionResult:
    """Result from SOL content detection stage"""
    has_sol_content: bool
    confidence: str  # "high", "medium", "low"
    detected_subject: Optional[str] = None
    detected_grade: Optional[str] = None
    detected_standards_count: int = 0
    reasoning: str = ""
    extracted_content: str = ""  # Relevant SOL content to pass to transformation


@dataclass
class ProcessingResult:
    """Result from complete pipeline processing"""
    success: bool
    source_file: str
    output_file: Optional[str] = None
    detection_result: Optional[SOLDetectionResult] = None
    standards_extracted: int = 0
    error: Optional[str] = None


class SOLPipeline:
    """Two-stage SOL document processing pipeline"""

    def __init__(
        self,
        openai_api_key: Optional[str] = None,
        staging_dir: str = "./sol_staging"
    ):
        """Initialize the SOL processing pipeline

        Args:
            openai_api_key: OpenAI API key for LLM processing
            staging_dir: Directory for staging output JSON files
        """
        self.openai_api_key = openai_api_key or os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            raise ValueError("OpenAI API key required. Set OPENAI_API_KEY environment variable.")

        self.client = OpenAI(api_key=self.openai_api_key)
        self.staging_dir = Path(staging_dir)
        self.staging_dir.mkdir(exist_ok=True)

        logger.info(f"SOL Pipeline initialized. Staging directory: {self.staging_dir}")

    # ==================== FILE EXTRACTION ====================

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract raw text from .docx file"""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise

    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract raw text from .xlsx file"""
        try:
            workbook = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

                for row in sheet.iter_rows(values_only=True):
                    row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text_parts.append(row_text)

            return '\n'.join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {e}")
            raise

    def extract_text_from_python(self, file_path: str) -> str:
        """Extract raw text from .py file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Python files often contain SOL data as dictionaries or structured data
            return content

        except Exception as e:
            logger.error(f"Error extracting text from Python {file_path}: {e}")
            raise

    def extract_text_from_json(self, file_path: str) -> str:
        """Extract raw text from .json file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)

            # Convert JSON to formatted text for analysis
            return json.dumps(json_data, indent=2)

        except Exception as e:
            logger.error(f"Error extracting text from JSON {file_path}: {e}")
            raise

    def extract_document_text(self, file_path: str) -> str:
        """Extract text from supported file formats"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix in ['.docx', '.doc']:
            return self.extract_text_from_docx(str(file_path))
        elif suffix in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(str(file_path))
        elif suffix == '.py':
            return self.extract_text_from_python(str(file_path))
        elif suffix == '.json':
            return self.extract_text_from_json(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

    # ==================== STAGE 1: DETECTION ====================

    def create_detection_prompt(self, document_text: str, file_name: str) -> str:
        """Create prompt for SOL content detection"""
        return f"""
You are analyzing a document to determine if it contains Virginia Standards of Learning (SOL) content.

Document: {file_name}
Content Preview (first 3000 chars):
{document_text[:3000]}

Your task is to analyze this document and determine:
1. Does this document contain Virginia SOL standards?
2. What subject area (mathematics, science, english, social studies, etc.)?
3. What grade level(s) (K, 1, 2, 3, 4, 5, 6, 7, 8, or course names like "Algebra 1", "Biology")?
4. How many distinct SOL standards appear to be present?
5. Extract ONLY the SOL-relevant content (standards, descriptions, objectives)

Respond in this EXACT JSON format:
{{
  "has_sol_content": true/false,
  "confidence": "high/medium/low",
  "detected_subject": "mathematics" or "science" or "english" etc. (null if no SOL),
  "detected_grade": "3" or "K" or "Algebra 1" etc. (null if no SOL),
  "detected_standards_count": number,
  "reasoning": "Brief explanation of your determination",
  "extracted_content": "Only the SOL standards and descriptions, excluding instructional guidance, examples, or non-standard text"
}}

Look for indicators like:
- Standard codes (e.g., "3.1", "K.2", "ALG.1", "BIO.3")
- Phrases like "The student will..."
- Content strands (Number Sense, Geometry, Life Processes, etc.)
- Grade-level organization
- Virginia SOL formatting patterns

CRITICAL: Return ONLY valid JSON, no markdown formatting, no additional text.
"""

    def detect_sol_content(self, document_text: str, file_name: str) -> SOLDetectionResult:
        """Stage 1: Detect if document contains SOL content

        Args:
            document_text: Raw text from document
            file_name: Name of source file

        Returns:
            SOLDetectionResult with detection findings
        """
        logger.info(f"🔍 Stage 1: Detecting SOL content in {file_name}")

        prompt = self.create_detection_prompt(document_text, file_name)

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert analyzer of Virginia Standards of Learning documents. Your job is to detect SOL content and extract only the relevant standards information."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_tokens=4000
            )

            response_text = response.choices[0].message.content.strip()

            # Clean markdown formatting if present
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()

            # Parse JSON response
            result_data = json.loads(response_text)

            detection_result = SOLDetectionResult(
                has_sol_content=result_data.get('has_sol_content', False),
                confidence=result_data.get('confidence', 'low'),
                detected_subject=result_data.get('detected_subject'),
                detected_grade=result_data.get('detected_grade'),
                detected_standards_count=result_data.get('detected_standards_count', 0),
                reasoning=result_data.get('reasoning', ''),
                extracted_content=result_data.get('extracted_content', '')
            )

            logger.info(
                f"✓ Detection complete: "
                f"has_SOL={detection_result.has_sol_content}, "
                f"confidence={detection_result.confidence}, "
                f"subject={detection_result.detected_subject}, "
                f"grade={detection_result.detected_grade}, "
                f"count={detection_result.detected_standards_count}"
            )

            return detection_result

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse detection JSON response: {e}")
            logger.error(f"Response was: {response_text[:500]}")
            raise
        except Exception as e:
            logger.error(f"Error during SOL detection: {e}")
            raise

    # ==================== STAGE 2: TRANSFORMATION ====================

    def create_transformation_prompt(self, sol_content: str, subject: str, grade: str) -> str:
        """Create prompt for transforming SOL content to standards.json format"""
        return f"""
You are transforming Virginia SOL standards into a specific JSON format.

Subject: {subject}
Grade: {grade}

SOL Content to Transform:
{sol_content}

Transform this content into the EXACT format below. The output must match this structure precisely:

{{
  "{subject}": {{
    "{grade}": {{
      "STANDARD_CODE": {{
        "title": "Brief title summarizing the standard",
        "description": "Full description starting with 'The student will...'",
        "strands": ["Content Strand Name"]
      }}
    }}
  }}
}}

EXAMPLE of correct format:
{{
  "mathematics": {{
    "3": {{
      "3.1": {{
        "title": "Read and represent whole numbers through 9,999",
        "description": "The student will read and write six-digit whole numbers and identify the place value and value of each digit.",
        "strands": ["Number and Number Sense"]
      }},
      "3.2": {{
        "title": "Compare and order whole numbers",
        "description": "The student will compare and order whole numbers through 9,999.",
        "strands": ["Number and Number Sense"]
      }}
    }}
  }}
}}

CRITICAL FORMATTING RULES:
1. Use the exact subject name provided: "{subject}"
2. Use the exact grade provided: "{grade}"
3. Standard codes should match the SOL format (e.g., "3.1", "K.2", "ALG.1")
4. Title should be concise (5-15 words)
5. Description should start with "The student will..." when possible
6. Strands are arrays and should use proper content strand names like:
   - Mathematics: "Number and Number Sense", "Computation and Estimation", "Measurement and Geometry", "Probability and Statistics", "Patterns, Functions, and Algebra"
   - Science: "Scientific Investigation, Reasoning, and Logic", "Force, Motion, and Energy", "Matter", "Life Processes", "Living Systems", "Interrelationships in Earth/Space Systems", "Earth Patterns, Cycles, and Change", "Earth Resources"
   - English: "Oral Language", "Reading", "Writing", "Research"
7. Extract ALL standards from the content
8. Return ONLY the JSON object - no markdown, no explanations, no extra text

Your response must start with {{ and end with }}
"""

    def transform_to_standards_format(
        self,
        sol_content: str,
        subject: str,
        grade: str
    ) -> Dict[str, Any]:
        """Stage 2: Transform SOL content to standards.json format

        Args:
            sol_content: Extracted SOL content from detection stage
            subject: Subject area (e.g., "mathematics")
            grade: Grade level (e.g., "3", "K")

        Returns:
            Dictionary in standards.json format
        """
        logger.info(f"🔄 Stage 2: Transforming SOL content to standards.json format")

        prompt = self.create_transformation_prompt(sol_content, subject, grade)

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert at transforming Virginia SOL standards into structured JSON format. You must follow the exact schema provided and return valid JSON only."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_tokens=4000
            )

            response_text = response.choices[0].message.content.strip()

            # Clean markdown formatting if present
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            response_text = response_text.strip()

            # Parse JSON response
            standards_data = json.loads(response_text)

            # Validate structure
            if not isinstance(standards_data, dict):
                raise ValueError("Transformed data is not a dictionary")

            # Count standards
            total_standards = 0
            for subj in standards_data.values():
                for grade_data in subj.values():
                    total_standards += len(grade_data)

            logger.info(f"✓ Transformation complete: {total_standards} standards extracted")

            return standards_data

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse transformation JSON response: {e}")
            logger.error(f"Response was: {response_text[:500]}")
            raise
        except Exception as e:
            logger.error(f"Error during SOL transformation: {e}")
            raise

    # ==================== PIPELINE ORCHESTRATION ====================

    def process_file(
        self,
        file_path: str,
        skip_detection: bool = False,
        force_subject: Optional[str] = None,
        force_grade: Optional[str] = None
    ) -> ProcessingResult:
        """Process a single file through the complete pipeline

        Args:
            file_path: Path to the document file (.docx or .xlsx)
            skip_detection: Skip detection stage if subject/grade are known
            force_subject: Override detected subject
            force_grade: Override detected grade

        Returns:
            ProcessingResult with processing outcome
        """
        file_path = Path(file_path)
        logger.info(f"\n{'='*60}")
        logger.info(f"📁 Processing: {file_path.name}")
        logger.info(f"{'='*60}")

        try:
            # Extract document text
            logger.info("📄 Extracting document text...")
            document_text = self.extract_document_text(str(file_path))
            logger.info(f"✓ Extracted {len(document_text)} characters")

            # Stage 1: Detection (unless skipped)
            if skip_detection and force_subject and force_grade:
                logger.info("⏭️  Skipping detection stage (forced parameters)")
                detection_result = SOLDetectionResult(
                    has_sol_content=True,
                    confidence="forced",
                    detected_subject=force_subject,
                    detected_grade=force_grade,
                    detected_standards_count=0,
                    reasoning="Detection skipped - using forced parameters",
                    extracted_content=document_text
                )
            else:
                detection_result = self.detect_sol_content(document_text, file_path.name)

            # Check if SOL content was detected
            if not detection_result.has_sol_content:
                logger.warning(f"❌ No SOL content detected in {file_path.name}")
                logger.warning(f"Reasoning: {detection_result.reasoning}")
                return ProcessingResult(
                    success=False,
                    source_file=str(file_path),
                    detection_result=detection_result,
                    error="No SOL content detected"
                )

            # Use forced parameters if provided, otherwise use detected
            subject = force_subject or detection_result.detected_subject
            grade = force_grade or detection_result.detected_grade

            if not subject or not grade:
                raise ValueError("Could not determine subject or grade level")

            # Stage 2: Transformation
            standards_data = self.transform_to_standards_format(
                detection_result.extracted_content,
                subject,
                grade
            )

            # Save to staging directory
            output_filename = f"{file_path.stem}_standards.json"
            output_path = self.staging_dir / output_filename

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(standards_data, f, indent=2, ensure_ascii=False)

            logger.info(f"💾 Saved to: {output_path}")

            # Count standards
            standards_count = 0
            for subj_data in standards_data.values():
                for grade_data in subj_data.values():
                    standards_count += len(grade_data)

            logger.info(f"✅ Processing complete: {standards_count} standards extracted")

            return ProcessingResult(
                success=True,
                source_file=str(file_path),
                output_file=str(output_path),
                detection_result=detection_result,
                standards_extracted=standards_count
            )

        except Exception as e:
            logger.error(f"❌ Processing failed for {file_path.name}: {e}")
            return ProcessingResult(
                success=False,
                source_file=str(file_path),
                error=str(e)
            )

    def get_supported_files(self, directory_path: str) -> List[Path]:
        """Get all supported files from a directory

        Args:
            directory_path: Path to directory containing SOL documents

        Returns:
            List of Path objects for all supported files
        """
        directory = Path(directory_path)
        supported_extensions = ['.docx', '.doc', '.xlsx', '.xls', '.py', '.json']
        all_files = []

        for ext in supported_extensions:
            # Use glob to find all files with this extension
            all_files.extend(directory.glob(f"*{ext}"))

        # Sort by name for consistent processing order
        return sorted(all_files)

    def process_directory(
        self,
        directory_path: str,
        patterns: List[str] = None,
        skip_detection: bool = False
    ) -> List[ProcessingResult]:
        """Process all matching files in a directory

        Args:
            directory_path: Path to directory containing SOL documents
            patterns: File patterns to match (default: auto-detect all supported types)
            skip_detection: Skip detection for all files

        Returns:
            List of ProcessingResult for each file
        """
        directory = Path(directory_path)

        if patterns is None:
            # Auto-detect all supported file types
            all_files = self.get_supported_files(directory_path)
            pattern_desc = "all supported files (.docx, .xlsx, .py, .json)"
        else:
            # Use specified patterns
            all_files = []
            for pattern in patterns:
                all_files.extend(directory.glob(pattern))
            all_files = sorted(all_files)
            pattern_desc = f"patterns: {patterns}"

        logger.info(f"\n{'='*60}")
        logger.info(f"📂 Processing directory: {directory}")
        logger.info(f"📋 Found {len(all_files)} files matching {pattern_desc}")
        logger.info(f"{'='*60}\n")

        results = []

        for i, file_path in enumerate(all_files, 1):
            logger.info(f"\n[{i}/{len(all_files)}] Processing {file_path.name}")

            result = self.process_file(
                str(file_path),
                skip_detection=skip_detection
            )
            results.append(result)

            # Brief delay to avoid rate limits
            if i < len(all_files):
                import time
                time.sleep(1)

        # Summary
        successful = sum(1 for r in results if r.success)
        total_standards = sum(r.standards_extracted for r in results if r.success)

        logger.info(f"\n{'='*60}")
        logger.info(f"📊 PROCESSING SUMMARY")
        logger.info(f"{'='*60}")
        logger.info(f"Total files processed: {len(results)}")
        logger.info(f"Successful: {successful}")
        logger.info(f"Failed: {len(results) - successful}")
        logger.info(f"Total standards extracted: {total_standards}")
        logger.info(f"Output directory: {self.staging_dir}")
        logger.info(f"{'='*60}\n")

        return results


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("""
SOL Processing Pipeline
Usage:
  python sol_pipeline.py <file_path>                    # Process single file
  python sol_pipeline.py <directory_path> --dir         # Process directory
  python sol_pipeline.py <file_path> --subject math --grade 3  # Force parameters
        """)
        sys.exit(1)

    path = sys.argv[1]
    is_dir = "--dir" in sys.argv

    # Parse force parameters
    force_subject = None
    force_grade = None
    if "--subject" in sys.argv:
        idx = sys.argv.index("--subject")
        force_subject = sys.argv[idx + 1] if idx + 1 < len(sys.argv) else None
    if "--grade" in sys.argv:
        idx = sys.argv.index("--grade")
        force_grade = sys.argv[idx + 1] if idx + 1 < len(sys.argv) else None

    try:
        pipeline = SOLPipeline()

        if is_dir:
            results = pipeline.process_directory(path)
        else:
            result = pipeline.process_file(
                path,
                skip_detection=(force_subject and force_grade),
                force_subject=force_subject,
                force_grade=force_grade
            )
            if result.success:
                print(f"\n✅ Success! Output saved to: {result.output_file}")
            else:
                print(f"\n❌ Failed: {result.error}")
                sys.exit(1)

    except Exception as e:
        print(f"\n❌ Pipeline error: {e}")
        sys.exit(1)
